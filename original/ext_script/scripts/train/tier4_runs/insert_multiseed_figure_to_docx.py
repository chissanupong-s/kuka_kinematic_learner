#!/usr/bin/env python3
"""Insert the multi-seed reproducibility figure (Fig 5.7) into the docx after the §5.4 fourth-observation paragraph."""
import shutil, sys
from copy import deepcopy
from pathlib import Path

import docx
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TIMESTAMP = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")
SRC = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v5_variance_{TIMESTAMP}.docx"
DST = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v6_final_{TIMESTAMP}.docx"
PNG = TRAIN_DIR / f"report_resources/figures/appendix_tb_plots/{TIMESTAMP}/fig_multiseed_reproducibility_{TIMESTAMP}.png"


def main():
    if not SRC.exists():
        sys.exit(f"Source not found: {SRC}")
    if not PNG.exists():
        sys.exit(f"Figure not found: {PNG}")
    shutil.copy(str(SRC), str(DST))
    doc = docx.Document(str(DST))

    # Find the §5.4 multi-seed paragraph just inserted (begins "A fourth observation comes from...")
    target = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("A fourth observation comes from the multi-seed evaluation"):
            target = p; break
    if target is None:
        sys.exit("Could not find §5.4 multi-seed paragraph anchor.")

    # New empty paragraph for the image, after the anchor
    img_p_el = deepcopy(target._element)
    for r in list(img_p_el.findall(qn('w:r'))):
        img_p_el.remove(r)
    target._element.addnext(img_p_el)
    img_p = Paragraph(img_p_el, target._parent)
    img_p.text = ""
    img_p.add_run().add_picture(str(PNG), width=Inches(6.0))

    # Caption after image
    cap_el = deepcopy(target._element)
    for r in list(cap_el.findall(qn('w:r'))):
        cap_el.remove(r)
    img_p_el.addnext(cap_el)
    cap_p = Paragraph(cap_el, target._parent)
    cap_p.text = (
        "Figure 5.7  Multi-seed reproducibility of the per-DoF adaptation result. "
        "Three seeds (42, 1, 2) hold the Stage-2 shared meta-kinematics checkpoint and "
        "Stage-3 hyperparameters fixed and vary only parameter initialisation, support/query "
        "split and minibatch sampling. Diamonds with whiskers show mean ± standard deviation (n = 3); "
        "individual seed results are plotted as filled markers. The dashed red line marks the single-task "
        "baseline for the corresponding DoF. The 5-DoF and 6-DoF results are reproducible to within "
        "the marker thickness; the 7-DoF result has notable seed sensitivity, with seed 42 attaining a "
        "lower error than seeds 1 and 2."
    )

    doc.save(str(DST))
    print(f"[OK] Saved: {DST}")


if __name__ == "__main__":
    main()
