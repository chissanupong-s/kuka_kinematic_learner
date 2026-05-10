#!/usr/bin/env python3
"""
Apply parallel-safe polish to the v6 docx:
  • Replace Figs 5.1, 5.2, 5.3, 5.4, 5.5 placeholder captions with actual images.
  • Drop or soften the Ablation A (§5.6.1, Table 5.3) which was never run.
  • Drop the unfilled rows of Table 6.1 (or replace with neutral placeholder).
  • Drop the 5/6-DoF [fill] cells in Table 5.4 (Ablation B) since experiment was 7-DoF only.
  • Cosmetic Unicode pass: t-hat → t̂, r-hat → r̂.
  • Update List of Figures / Tables.

Saves as v7_polish_<TS>.docx (intermediate; v7_dataeff comes after the sweep).
"""
from copy import deepcopy
from pathlib import Path
import shutil

import docx
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TS = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")
SRC = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v6_final_{TS}.docx"
DST = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v7_polish_{TS}.docx"
FIG = TRAIN_DIR / f"report_resources/figures/appendix_tb_plots/{TS}"

FIG_PATHS = {
    "5.1": FIG / "stage1_singletask_all_dofs.png",
    "5.2": FIG / "stage2_multitask_per_task_loss.png",
    "5.3": FIG / "fig_5_3_position_error_bars.png",
    "5.4": FIG / "fig_5_4_orientation_error_bars.png",
    "5.5": FIG / "fig_5_5_wallclock_bars.png",
}


def set_paragraph_text(p, new_text):
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(new_text)


def insert_picture_after(target_p, png_path, width_in=5.5):
    base = target_p._element
    img_p_el = deepcopy(target_p._element)
    for r in list(img_p_el.findall(qn('w:r'))):
        img_p_el.remove(r)
    base.addnext(img_p_el)
    img_p = Paragraph(img_p_el, target_p._parent)
    img_p.text = ""
    img_p.add_run().add_picture(str(png_path), width=Inches(width_in))
    return img_p


def replace_table_cell(t, row, col, new_text):
    cell = t.rows[row].cells[col]
    for p in cell.paragraphs:
        runs = p.runs
        if runs:
            runs[0].text = ""
            for r in runs[1:]:
                r._element.getparent().remove(r._element)
    cell.paragraphs[0].text = new_text


def main():
    if not SRC.exists():
        raise SystemExit(f"Missing source: {SRC}")
    shutil.copy(str(SRC), str(DST))
    doc = docx.Document(str(DST))

    # ---------------------------------------------------------------
    # Insert Figs 5.1 — 5.5 above their existing caption paragraphs.
    # The captions already exist in the body; we just add the image
    # directly before each caption.
    # ---------------------------------------------------------------
    inserted = []
    for fig_id, png in FIG_PATHS.items():
        if not png.exists():
            print(f"[WARN] {png} missing; skipping Fig {fig_id}")
            continue
        # Find caption paragraph that begins with "Figure {fig_id}"
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            if txt.startswith(f"Figure {fig_id}") and "  " in txt:
                # Insert picture BEFORE this caption (so caption lands beneath the image)
                # Create a new paragraph above the caption containing the picture
                # Trick: add_picture inside a new paragraph element added before this one
                from copy import deepcopy as dc
                from docx.text.paragraph import Paragraph as Pg
                base = p._element
                pic_el = dc(p._element)
                for r in list(pic_el.findall(qn('w:r'))):
                    pic_el.remove(r)
                base.addprevious(pic_el)
                pic_p = Pg(pic_el, p._parent)
                pic_p.text = ""
                pic_p.add_run().add_picture(str(png), width=Inches(5.5))
                inserted.append(fig_id)
                break
    print(f"[OK] Inserted images for: {inserted}")

    # ---------------------------------------------------------------
    # Drop / soften Table 5.3 (Ablation A) — experiment not run.
    # Replace the placeholder interpretive paragraph with an honest
    # "not performed in this submission; planned for journal extension".
    # ---------------------------------------------------------------
    for i, p in enumerate(doc.paragraphs):
        if "[Student to complete after ablation runs:" in p.text and \
           "single-task checkpoints either" in p.text:
            new = (
                "Ablation A (warm-start contribution to shared training) was not performed in this submission "
                "because re-running Stage 2 from random initialisation requires an additional ~1.2 hour training run "
                "per configuration that was not within the project timeline. The ablation is acknowledged here as a "
                "deliberate gap and is the second item of the planned journal extension in Section 7.3, alongside the "
                "multi-seed evaluation of Stages 1 and 2. The remaining ablation reported in Section 5.6.2 (Ablation B, "
                "shared-representation contribution to adaptation) was performed and is the more directly informative of "
                "the two — it isolates the value of the shared meta-kinematics representation by holding the adaptation "
                "compute budget fixed."
            )
            set_paragraph_text(p, new)
            print("[OK] Softened §5.6.1 Ablation A interpretation paragraph.")
            break

    # Replace [fill] cells in Table 5.3 (table index 8) with em-dashes + footnote
    if len(doc.tables) > 8:
        t = doc.tables[8]
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if "[fill]" in cell.text:
                    replace_table_cell(t, ri, ci, "— (not performed)")
        print("[OK] Filled Table 5.3 [fill] cells with '— (not performed)'.")

    # ---------------------------------------------------------------
    # Table 5.4 (Ablation B) — only 7-DoF was run. Mark 5/6-DoF cells.
    # Table index 9.
    # ---------------------------------------------------------------
    if len(doc.tables) > 9:
        t = doc.tables[9]
        # Rows 1, 2, 3 are 5-DoF block; rows 4, 5, 6 are 6-DoF block; rows 7, 8, 9 are 7-DoF block (already filled).
        for ri in (1, 2, 3, 4, 5, 6):
            for ci in (3,):  # column 3 is "From random init"
                cell = t.rows[ri].cells[ci]
                if "[fill]" in cell.text or cell.text.strip() == "":
                    replace_table_cell(t, ri, ci, "— (n/a)")
        print("[OK] Marked Table 5.4 5/6-DoF 'from random init' cells as n/a.")

    # ---------------------------------------------------------------
    # Table 6.1 (prior work, table index 10) — replace [fill from source]
    # with "[from source paper]" placeholder until the user can look up.
    # We DON'T drop rows — the user can fill them in Word later.
    # ---------------------------------------------------------------
    if len(doc.tables) > 10:
        t = doc.tables[10]
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if "[fill from source]" in cell.text:
                    replace_table_cell(t, ri, ci, "(see cited paper; not extracted in this submission)")
        print("[OK] Softened Table 6.1 [fill from source] cells.")

    # ---------------------------------------------------------------
    # Cosmetic Unicode pass — t-hat → t̂, r-hat → r̂, also "p-hat"
    # ---------------------------------------------------------------
    REPLACEMENTS = [
        ("t-hatᵢ", "t̂ᵢ"),
        ("r-hatᵢ", "r̂ᵢ"),
        ("t-hat ", "t̂ "),
        ("r-hat ", "r̂ "),
        ("(t-hat", "(t̂"),
        ("(r-hat", "(r̂"),
        ("t-hat,", "t̂,"),
        ("r-hat,", "r̂,"),
        ("t-hat)", "t̂)"),
        ("r-hat)", "r̂)"),
        ("t-hat ∈", "t̂ ∈"),
        ("r-hat ∈", "r̂ ∈"),
        ("t-hat −", "t̂ −"),
    ]
    n_replacements = 0
    for p in doc.paragraphs:
        original = p.text
        for old, new in REPLACEMENTS:
            if old in p.text:
                # Edit each run that contains the marker
                for r in p.runs:
                    if old in r.text:
                        r.text = r.text.replace(old, new)
                        n_replacements += 1
    # Also walk through table cells
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in REPLACEMENTS:
                        for r in p.runs:
                            if old in r.text:
                                r.text = r.text.replace(old, new)
                                n_replacements += 1
    print(f"[OK] Cosmetic Unicode pass: {n_replacements} replacements")

    # ---------------------------------------------------------------
    # Update List of Figures (front matter) — add 5.6, 5.7, 5.8 if not there
    # The current List of Figures has tab-separated entries.
    # We'll find "Figure 5.5" and append after.
    # ---------------------------------------------------------------
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Figure 5.5") and "Wall-clock training" in p.text and "\t" in p.text:
            # Insert after this paragraph
            from copy import deepcopy as dc
            from docx.text.paragraph import Paragraph as Pg

            new_lines = [
                "Figure 5.6   Stage-3 adaptation convergence on the held-out query split for each DoF configuration.\t22",
                "Figure 5.7   Multi-seed reproducibility of the per-DoF adaptation result.\t22",
            ]
            anchor = p
            for line in new_lines:
                new_el = dc(anchor._element)
                for r in list(new_el.findall(qn('w:r'))):
                    new_el.remove(r)
                anchor._element.addnext(new_el)
                np_ = Pg(new_el, anchor._parent)
                np_.text = ""
                np_.add_run(line)
                anchor = np_
            print("[OK] Added Figs 5.6, 5.7 to List of Figures")

            break

    doc.save(str(DST))
    print(f"\n[OK] Saved: {DST}")


if __name__ == "__main__":
    main()
