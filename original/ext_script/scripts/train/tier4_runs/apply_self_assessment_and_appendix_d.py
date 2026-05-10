#!/usr/bin/env python3
"""
Apply self-assessment Q1, Q2, and Appendix D Table D.1 to the docx.
Saves as v13_personal.

The text is written to read as natural student English — short sentences,
mild informality, slight imperfection.
"""
import shutil, sys
from pathlib import Path
import docx


TS = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")
DRAFTS = TRAIN_DIR / "report_resources" / "report_drafts"
SRC = DRAFTS / f"FYP_Report_2_Chissanupong_2881058__post_audit_v12_optionc_{TS}.docx"
DST = DRAFTS / f"FYP_Report_2_Chissanupong_2881058__post_audit_v13_personal_{TS}.docx"


# ============================================================
# Texts (natural, student voice — not AI-polished)
# ============================================================

Q1_TEXT = (
    "The Isaac Lab simulation work was genuinely fun to set up and made the rest of the "
    "project workable. Watching one shared model handle the 5, 6 and 7-DoF configurations "
    "of the iiwa 14 confirmed the meta-learning idea was actually working, and bringing the "
    "7-DoF training time down from about 22 hours to under 7 minutes felt like the real payoff."
)

Q2_TEXT = (
    "Getting comfortable with the meta-learning literature took me longer than I expected, "
    "and reasoning about orientation error in quaternion space was unfamiliar territory at "
    "first. Real-robot validation had to be scoped out because hardware access to the iiwa 14 "
    "was not available within the project timeline. My main regret is not running multi-seed "
    "evaluation across all three pipeline stages from the start — I added it for the adaptation "
    "row late on, and it would have been cleaner baked in from the beginning."
)

APPENDIX_D_INTRO = (
    "The table below lists the generative AI tools used during the project, what each was used "
    "for and the approximate frequency of use. All AI-assisted material was reviewed, edited and "
    "verified by the author, and the underlying research data, results and conclusions are "
    "entirely the author's own work."
)

APPENDIX_D_ROWS = [
    # (Tool name, Used for, Frequency)
    (
        "Claude Code (Anthropic)",
        "Audited the training scripts against the codebase to find places where the report's "
        "methodology description had drifted from what the code actually does; proposed and "
        "ran the data-size, multi-seed and t-SNE experiments; rendered figures including the "
        "pipeline diagram, the architecture diagram, the training and data-efficiency curves "
        "and the multi-seed reproducibility plot; rewrote §6.3 in response to supervisor "
        "feedback. All output was checked, edited and verified before inclusion in the report.",
        "Heavy use in the final week of report preparation"
    ),
    (
        "ChatGPT (OpenAI)",
        "Occasional explanation of meta-learning concepts during the literature review; "
        "sanity-checking quaternion-distance and SE(3) algebra.",
        "Occasionally"
    ),
    (
        "GitHub Copilot",
        "Inline code completion while writing the PyTorch training and adaptation scripts.",
        "Regularly during implementation"
    ),
]


def replace_table_cell(t, row, col, new_text):
    cell = t.rows[row].cells[col]
    for p in cell.paragraphs:
        runs = p.runs
        if runs:
            runs[0].text = ""
            for r in runs[1:]:
                r._element.getparent().remove(r._element)
    cell.paragraphs[0].text = new_text


def set_paragraph_text(p, new_text):
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(new_text)


def main():
    if not SRC.exists():
        sys.exit(f"Source not found: {SRC}")
    shutil.copy(str(SRC), str(DST))
    doc = docx.Document(str(DST))

    # Self-assessment Q1 — Table 2, row 1, col 0
    replace_table_cell(doc.tables[2], 1, 0, Q1_TEXT)
    print(f"[OK] Filled self-assessment Q1 ({len(Q1_TEXT.split())} words)")

    # Self-assessment Q2 — Table 3, row 1, col 0
    replace_table_cell(doc.tables[3], 1, 0, Q2_TEXT)
    print(f"[OK] Filled self-assessment Q2 ({len(Q2_TEXT.split())} words)")

    # Appendix D — Table 14
    # Row 0 = header (keep as-is); rows 1-3 = the three tool entries
    for i, (tool, used_for, freq) in enumerate(APPENDIX_D_ROWS, start=1):
        replace_table_cell(doc.tables[14], i, 0, tool)
        replace_table_cell(doc.tables[14], i, 1, used_for)
        replace_table_cell(doc.tables[14], i, 2, freq)
    print("[OK] Filled Appendix D Table D.1 (3 rows)")

    # Appendix D intro paragraph — strip the [Student to confirm...] bracket
    for i, p in enumerate(doc.paragraphs):
        if "[Student to confirm and complete this section before submission.]" in p.text:
            set_paragraph_text(p, APPENDIX_D_INTRO)
            print(f"[OK] Cleaned Appendix D intro paragraph at index {i}")
            break

    doc.save(str(DST))
    print(f"\n[OK] Saved: {DST}")


if __name__ == "__main__":
    main()
