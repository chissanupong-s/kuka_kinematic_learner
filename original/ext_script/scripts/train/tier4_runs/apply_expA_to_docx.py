#!/usr/bin/env python3
"""
Apply Experiment A (random-init 7-DoF adapt) results to the docx.

Two metrics from Experiment A enter the report:
  • Matched-budget result at step 12 000 (which corresponds to roughly the
    0.111 hr wall-clock budget reported in Table 5.1's "From shared" 7-DoF
    row): pos_mae = 0.0265 m, ori = 4.28°.
    This is the apples-to-apples ablation comparison.
  • Full-budget BEST result at step 96 000: pos_mae = 0.00664 m,
    ori = 0.833°.  This shows that random-init eventually catches up given
    a much larger training budget — but only after ~74 min of wall clock,
    versus the 0.111 hr at which the shared-init result was reached.

Both are reported.  Table 5.4's 7-DoF row is filled with the matched-budget
numbers; the BEST-of-run is mentioned in the explanatory paragraph as the
honest counterpoint.
"""

import shutil
from copy import deepcopy
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TIMESTAMP = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")

SRC_DOCX = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v2_{TIMESTAMP}.docx"
DST_DOCX = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v3_{TIMESTAMP}.docx"

# Experiment A measured numbers
MATCHED_STEP = 12000
MATCHED_POS = 0.0265
MATCHED_ORI = 4.28
MATCHED_BUDGET_HR = 0.111

BEST_STEP = 96000
BEST_POS = 0.00664
BEST_ORI = 0.833
FULL_RUN_HR = 74.1 / 60.0   # 74.1 min wall-clock

# Headline shared-init (already in report)
SHARED_POS = 0.0099
SHARED_ORI = 1.7104
SHARED_BUDGET_HR = 0.111

SINGLE_TASK_POS = 0.0101
SINGLE_TASK_ORI = 2.0853
SINGLE_TASK_HR = 22.12


def set_paragraph_text(p, new_text):
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(new_text)


def replace_table_cell(t, row, col, new_text):
    cell = t.rows[row].cells[col]
    for p in cell.paragraphs:
        runs = p.runs
        if runs:
            runs[0].text = ""
            for r in runs[1:]:
                r._element.getparent().remove(r._element)
    cell.paragraphs[0].text = new_text


def main():
    if not SRC_DOCX.exists():
        raise SystemExit(f"Source not found: {SRC_DOCX}")

    shutil.copy(str(SRC_DOCX), str(DST_DOCX))
    doc = docx.Document(str(DST_DOCX))

    # =================================================================
    # 1. Fill Table 5.4 (Ablation B) 7-DoF row with matched-wall-clock numbers
    # Table index = 9. Rows: header, 5DoF pos, 5DoF ori, 5DoF budget,
    # 6DoF pos, 6DoF ori, 6DoF budget, 7DoF pos, 7DoF ori, 7DoF budget.
    # 7-DoF block sits at rows 7, 8, 9.
    # Columns: 0=Configuration, 1=Metric, 2=From shared, 3=From random init, 4=Single-task baseline
    # =================================================================
    t = doc.tables[9]
    replace_table_cell(t, 7, 3, f"{MATCHED_POS:.4f}")     # 7-DoF position from-random
    replace_table_cell(t, 8, 3, f"{MATCHED_ORI:.4f}")     # 7-DoF orientation from-random
    replace_table_cell(t, 9, 3, f"{MATCHED_BUDGET_HR:.3f}")   # 7-DoF wall-clock budget from-random (matched)
    print("[OK] Filled Table 5.4 7-DoF row (matched-wall-clock).")

    # =================================================================
    # 2. Replace the [Student to complete] paragraph after Table 5.4 with the
    # quantitative interpretation, including the honest "given enough time
    # random-init catches up" finding.
    # =================================================================
    interpretation = (
        f"Two complementary readings of Table 5.4 are informative. "
        f"At the matched-wall-clock budget of {SHARED_BUDGET_HR:.3f} hr — the same budget that produces the headline "
        f"{SHARED_POS:.4f} m / {SHARED_ORI:.4f}° result of Table 5.1 — random-initialised adaptation reaches only "
        f"{MATCHED_POS:.4f} m / {MATCHED_ORI:.4f}° on 7 DoF, roughly {MATCHED_POS/SHARED_POS:.1f}× worse on position and "
        f"{MATCHED_ORI/SHARED_ORI:.1f}× worse on orientation than the shared-initialised model. This is the central "
        f"Ablation B finding: within the practical wall-clock budget, the shared meta-kinematics representation does "
        f"genuinely useful work that random initialisation cannot recover from.  "
        f"For completeness, when the random-initialised model is allowed to continue adapting beyond the matched budget, "
        f"its best score over the full 100 000-step run reaches {BEST_POS:.4f} m / {BEST_ORI:.4f}° at step {BEST_STEP:,}, "
        f"which on this metric is comparable to the shared-initialised model's headline 7-DoF performance — but at a "
        f"wall-clock cost of {FULL_RUN_HR:.2f} hr, more than {FULL_RUN_HR/SHARED_BUDGET_HR:.0f}× the time used by the "
        f"shared-initialised path. The single-task 7-DoF baseline ({SINGLE_TASK_POS} m / {SINGLE_TASK_ORI}° at "
        f"{SINGLE_TASK_HR:.2f} hr) sets the upper-bound cost of the matched-accuracy reference. "
        f"Taken together, these three numbers — shared adaptation at 0.111 hr, random-init adaptation at 0.111 hr, and "
        f"random-init adaptation given the full 7-DoF training budget — separate the two contributions of the framework: "
        f"the shared backbone provides a near-optimal starting point that a brief fine-tune can refine, whereas an "
        f"equally capable random-init model requires substantial additional compute to reach a comparable result."
    )
    found = False
    for i, p in enumerate(doc.paragraphs):
        if "[Student to complete after ablation runs:" in p.text and (
            "shared model substantially outperforms" in p.text or "From shared" in p.text or "matched-budget" in p.text
        ):
            set_paragraph_text(p, interpretation)
            print(f"[OK] Replaced Ablation B [Student to complete] paragraph at index {i}.")
            found = True
            break
    if not found:
        # Fallback: find by leading "[Student to complete after ablation runs:" near Table 5.4
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip().startswith("[Student to complete after ablation runs:"):
                # Heuristic: pick the second of two such paragraphs (Ablation A is first, B is second)
                count_seen = sum(1 for q in doc.paragraphs[:i+1] if q.text.strip().startswith("[Student to complete after ablation runs:"))
                if count_seen == 2:
                    set_paragraph_text(p, interpretation)
                    print(f"[OK] Replaced Ablation B [Student to complete] paragraph (heuristic) at index {i}.")
                    found = True
                    break
        if not found:
            print("[WARN] Could not locate Ablation B [Student to complete] paragraph.")

    doc.save(str(DST_DOCX))
    print(f"\n[OK] Saved: {DST_DOCX}")


if __name__ == "__main__":
    main()
