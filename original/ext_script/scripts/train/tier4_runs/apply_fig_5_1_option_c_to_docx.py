#!/usr/bin/env python3
"""
Apply the Option-C Figure 5.1 to the docx + add the explanatory sentence to §5.1.
Saves as v12_optionc.
"""
import shutil, sys
from copy import deepcopy
from pathlib import Path

import docx
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TS = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")
DRAFTS = TRAIN_DIR / "report_resources" / "report_drafts"
SRC = DRAFTS / f"FYP_Report_2_Chissanupong_2881058__post_audit_v11_dataeff_{TS}.docx"
DST = DRAFTS / f"FYP_Report_2_Chissanupong_2881058__post_audit_v12_optionc_{TS}.docx"
PNG_NEW = TRAIN_DIR / f"report_resources/figures/appendix_tb_plots/{TS}/fig_5_1_singletask_option_c.png"


def set_paragraph_text(p, new_text):
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(new_text)


def main():
    if not SRC.exists():
        sys.exit(f"Source missing: {SRC}")
    if not PNG_NEW.exists():
        sys.exit(f"New figure missing: {PNG_NEW}")

    shutil.copy(str(SRC), str(DST))
    doc = docx.Document(str(DST))

    # 1. Find Fig 5.1 caption (in body, not List of Figures)
    cap_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Figure 5.1") and "Single-task" in p.text and i > 150:
            cap_idx = i
            break
    if cap_idx is None:
        sys.exit("Fig 5.1 caption not found in body")
    cap_p = doc.paragraphs[cap_idx]

    # 2. The image (inserted in v7_polish) is the immediately previous paragraph that has a run with a drawing
    pic_p = doc.paragraphs[cap_idx - 1]

    # 3. Replace the picture: clear runs, add new picture
    runs = pic_p.runs
    for r in runs:
        r._element.getparent().remove(r._element)
    pic_p.add_run().add_picture(str(PNG_NEW), width=Inches(6.5))
    print("[OK] Replaced Fig 5.1 image with Option-C version")

    # 4. Update the Fig 5.1 caption to mention best-val markers
    new_cap = (
        "Figure 5.1  Single-task forward-kinematics training curves for the 5, 6 and 7 DoF configurations. "
        "Solid lines show training loss (smoothed exponential moving average; faded raw curves underneath); dashed lines show validation loss. "
        "Vertical dotted lines and markers indicate the per-DoF validation-loss minimum — the checkpoint saved automatically by the training script "
        "and used for the test-set evaluation reported in Table 5.1. The y-axis is logarithmic to make the train–validation gap visible; "
        "the gap after the validation minimum reflects the optimiser fitting training-set detail under a smaller learning rate, "
        "but this over-fit tail is observed only and is not deployed."
    )
    set_paragraph_text(cap_p, new_cap)
    print("[OK] Updated Fig 5.1 caption")

    # 5. Add the explanatory sentence to the §5.1 paragraph that introduces Fig 5.1.
    # The §5.1 paragraph begins "The single-task ResMLPs converged well…"
    explainer = (
        " The descending step-wise pattern in the training curves reflects the ReduceLROnPlateau scheduler "
        "halving the learning rate each time the validation loss stalls, allowing the optimiser progressively finer access to training-set detail. "
        "The plateau in validation loss — and the corresponding train–validation gap — indicates that the ResMLP has reached its effective generalisation "
        "ceiling for this task at the given capacity (1024-dimensional, 8 residual blocks); the test errors reported in Table 5.1 are computed from the "
        "validation-minimum checkpoint and therefore reflect this generalisation ceiling rather than the over-fit tail of the curve."
    )
    found = False
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("The single-task ResMLPs converged well"):
            # Append to existing paragraph
            last_run = p.runs[-1] if p.runs else p.add_run("")
            last_run.text = (last_run.text or "") + explainer
            print(f"[OK] Appended explainer to §5.1 paragraph at index {i}")
            found = True
            break
    if not found:
        print("[WARN] §5.1 anchor paragraph not found")

    doc.save(str(DST))
    print(f"\n[OK] Saved: {DST}")


if __name__ == "__main__":
    main()
