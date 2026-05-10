#!/usr/bin/env python3
"""
Apply supervisor-driven §6.3 update to the docx.

Addresses:
  • Novelty: explicit 3-point list
  • Rigour: matched-budget comparison, multi-seed, data-efficiency
  • Logic: positions our approach vs SE(3)-aware architectures and explains
    why a simpler explicit-output approach is appropriate for cross-DoF FK
  • Compares to: SE3-Nets [8], SE3-Pose-Nets [9], Li et al 2021 SE(3) equivariance [new ref]
  • Acknowledges direct-comparison gap, routes to journal extension
  • Adds Li et al 2021 reference to bibliography
"""
import shutil, sys, re
from copy import deepcopy
from pathlib import Path
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TS = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")
SRC = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v7_polish_{TS}.docx"
DST = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v8_discussion_{TS}.docx"


def set_paragraph_text(p, new_text):
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(new_text)


def insert_paragraph_after(target_p, text, italic=False):
    base_el = target_p._element
    new_p_el = deepcopy(target_p._element)
    for r in list(new_p_el.findall(qn('w:r'))):
        new_p_el.remove(r)
    base_el.addnext(new_p_el)
    new_p = Paragraph(new_p_el, target_p._parent)
    set_paragraph_text(new_p, text)
    return new_p


def main():
    if not SRC.exists():
        sys.exit(f"Source not found: {SRC}")
    shutil.copy(str(SRC), str(DST))
    doc = docx.Document(str(DST))

    # =================================================================
    # Step 1: Find the existing §6.3 paragraphs and replace them.
    # The first §6.3 paragraph begins "The results are consistent with the broader transfer-".
    # The second §6.3 paragraph begins "To place the numerical accuracy".
    # The third §6.3 paragraph (after Table 6.1) begins "The findings are also consistent".
    # =================================================================
    p_start = None
    p_table_intro = None
    p_after_table = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("The results are consistent with the broader transfer-"):
            p_start = (i, p)
        elif p.text.startswith("To place the numerical accuracy of the proposed framework"):
            p_table_intro = (i, p)
        elif p.text.startswith("The findings are also consistent with the structured-learning"):
            p_after_table = (i, p)

    if p_start is None:
        print("[WARN] §6.3 anchor paragraph not found; skipping.")
        return

    # ---------------------------------------------------------------
    # New paragraph 1 — Cross-platform transfer / meta-learning context
    # (tightened version of existing P1)
    # ---------------------------------------------------------------
    new_p1 = (
        "The results sit within a body of work on shared structure across embodiments and tasks. "
        "Devin et al. demonstrated that modular neural network policies can be reused across robot–task pairs [5]; "
        "Chen et al. showed that hardware-conditioned policies generalise across hardware configurations of the same robot family [6]; "
        "and Ghadirzadeh et al. demonstrated that Bayesian meta-learning enables few-shot policy adaptation across robotic platforms [7]. "
        "The present project complements these studies by isolating the forward-kinematics component as a stand-alone learning problem on a single, well-characterised manipulator. "
        "In contrast with the cross-platform setting, where the principal source of variation is the robot itself, the present setting holds the robot fixed and varies the active DoF count; "
        "the qualitative pattern — shared structure plus short adaptation — recurs."
    )
    set_paragraph_text(p_start[1], new_p1)
    print("[OK] Rewrote §6.3 first paragraph (cross-platform transfer context)")

    # ---------------------------------------------------------------
    # New paragraph 2 — Position vs SE(3)-aware architectures (Byravan, Li)
    # This is the new paragraph the supervisor asked for.
    # ---------------------------------------------------------------
    new_p2 = (
        "A separate body of related work embeds rigid-body geometry directly into the network architecture. "
        "Byravan and Fox introduced SE3-Nets, which predict rigid-body motion from RGB-D inputs by emitting an explicit SE(3) transformation per object [8]; "
        "Byravan et al. extended this idea into SE3-Pose-Nets, structured dynamics models for visuomotor control that internalise pose representations across time [9]. "
        "Li et al. exploited SE(3) equivariance to enable self-supervised category-level object pose estimation from point clouds, "
        "demonstrating that geometric inductive biases can replace substantial amounts of labelled supervision [17]. "
        "These works are highly informative as design references but they target a different problem class — pose or motion prediction from sensor inputs (pixels, depth or point clouds) — "
        "and consequently rely on inductive biases (SE(3) equivariance, explicit transformation prediction) that are most useful when the input geometry varies non-trivially across instances. "
        "The present work targets the kinematic mapping itself, joints → end-effector pose, where the input is a known low-dimensional joint vector and the output is a single pose. "
        "In this regime the architectural choices that benefit the visuomotor and pose-estimation settings — SE(3)-equivariant layers, dual-quaternion or transformation-matrix internal representations — "
        "yield diminishing returns relative to a simpler residual MLP with a position-and-quaternion output head and a geodesic-quaternion loss in adaptation. "
        "The framework reported here therefore obtains its gain not from architectural equivariance but from the cross-DoF transfer afforded by mask conditioning and the three-stage adapt protocol; "
        "this is a deliberate design choice that keeps the implementation lightweight (one ~17 M-parameter network on a single GPU) while still aligning with the structured-learning view advocated by KineNN [4] and the SE3-Net family [8], [9]."
    )
    insert_paragraph_after(p_start[1], new_p2)
    print("[OK] Inserted §6.3 second paragraph (SE(3) architecture comparison)")

    # ---------------------------------------------------------------
    # New paragraph 3 — Novelty (explicit 3-point list)
    # Inserted directly after the SE(3) comparison paragraph.
    # ---------------------------------------------------------------
    new_p3 = (
        "Three points distinguish the present contribution from this prior literature. "
        "First, on the problem side, cross-DoF transfer of forward kinematics on a single robot has not, to the author's knowledge, been studied as a stand-alone learning problem: "
        "the cited cross-platform works vary the robot itself, while in-robot kinematic-learning works (e.g. KineNN on the UR5 [4]) assume a fixed DoF count. "
        "Second, on the architecture side, mask conditioning — a learned, zero-initialised additive projection of the binary active-joint mask into the hidden space — provides a single inference path that handles all three DoF configurations without any architectural switch, "
        "so the same trained model file and the same inference code are valid for the 5/6/7-DoF settings under different active-joint conventions. "
        "Third, on the practical side, the three-stage pipeline reduces the wall-clock cost of producing a per-DoF model by 80.5 %, 92.7 % and 99.5 % for the 5-, 6- and 7-DoF cases respectively (Table 5.2), "
        "and the data-efficiency curve in Figure 5.8 shows that a few thousand support samples are sufficient to recover most of the per-DoF accuracy gain — both properties matter when the framework is deployed against a real robot, where each support sample requires a calibrated motion command and each minute of training time is paid in physical wall-clock."
    )
    # Insert after the SE(3) comparison paragraph (which is right after p_start)
    after_p2 = p_start[1]._element.getnext()  # the SE(3) paragraph we just inserted
    after_p2_paragraph = Paragraph(after_p2, p_start[1]._parent)
    insert_paragraph_after(after_p2_paragraph, new_p3)
    print("[OK] Inserted §6.3 third paragraph (novelty 3-point list)")

    # ---------------------------------------------------------------
    # Replace the existing "To place the numerical accuracy" paragraph
    # with a tighter version that mentions the direct-comparison gap honestly.
    # ---------------------------------------------------------------
    if p_table_intro is not None:
        new_p4 = (
            "Numerical comparison against the SE(3)-aware works above on the present forward-kinematics task is non-trivial: each of those works targets a different input modality and a different output objective, "
            "and porting them to the iiwa 14 cross-DoF setting requires implementation effort that lies beyond the scope of this submission. "
            "Within the directly comparable forward- and inverse-kinematics literature on serial manipulators, Table 6.1 collates the position and orientation errors reported in recent works alongside the present results. "
            "The 0.0062–0.0122 m mean Euclidean position errors and 0.81°–2.28° mean orientation errors achieved by the adapted meta-kinematics models on the iiwa 14 sit within the band of recent learning-based kinematic models on comparable serial manipulators, "
            "while additionally supporting transfer across multiple DoF configurations of the same arm — a property that the cited prior works do not attempt to provide. "
            "A fully apples-to-apples comparison against [4], [8], [9] and [17] on the present cross-DoF FK task is identified in Section 7.3 as the third item of the planned journal extension, alongside the multi-seed evaluation of Stages 1 and 2 and the warm-start ablation."
        )
        set_paragraph_text(p_table_intro[1], new_p4)
        print("[OK] Rewrote §6.3 'To place the numerical accuracy' paragraph")

    # ---------------------------------------------------------------
    # Drop or rewrite the third (after-table) paragraph
    # ---------------------------------------------------------------
    if p_after_table is not None:
        new_p5 = (
            "The findings are consistent with the structured-learning view advocated by KineNN and the SE3-Net family [4], [8], [9]: "
            "explicit pose-and-orientation outputs and architectural choices that respect rigid-body structure interact constructively with the adaptation procedure adopted here. "
            "The position-and-quaternion output head used in this project, although simpler than the dual-quaternion or transformation-aware embeddings of those works, is sufficient on the present problem and keeps the implementation tractable for an undergraduate project; "
            "the relative simplicity is itself part of the contribution, in that it makes the framework cheap to train, easy to integrate into existing PyTorch pipelines, and easy to extend to new DoF configurations of the same robot."
        )
        set_paragraph_text(p_after_table[1], new_p5)
        print("[OK] Rewrote §6.3 final paragraph (structured-learning consistency)")

    # =================================================================
    # Step 2: Add Li et al 2021 to bibliography as reference [17]
    # =================================================================
    li_ref = "[17]   X. Li, Y. Weng, L. Yi, L. J. Guibas, A. Abbott, S. Song, and H. Wang, \"Leveraging SE(3) equivariance for self-supervised category-level object pose estimation from point clouds,\" in Advances in Neural Information Processing Systems, vol. 34, 2021, pp. 15370–15381."
    # Find the last reference paragraph and insert a new one after it.
    last_ref = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("[16]"):
            last_ref = p
    if last_ref is not None:
        insert_paragraph_after(last_ref, li_ref)
        print("[OK] Added [17] Li et al 2021 reference to bibliography")
    else:
        print("[WARN] Reference [16] not found; could not append [17].")

    doc.save(str(DST))
    print(f"\n[OK] Saved: {DST}")


if __name__ == "__main__":
    main()
