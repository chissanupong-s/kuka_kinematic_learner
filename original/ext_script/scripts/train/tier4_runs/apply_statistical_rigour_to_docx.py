#!/usr/bin/env python3
"""
Add statistical rigour to v8 → v9_stats:
  • Compute 95% CI (t-distribution, df=2) for each DoF Adapted (best) row
  • Compute one-sample t-test against single-task baseline for each (DoF, metric)
  • Compute one-sample t-test against shared meta-kinematics baseline
  • Render Fig 5.7-CI: per-seed scatter + mean + t-CI shading (replaces old Fig 5.7)
  • Update Table 5.1 captions and the §5.4 fourth-observation paragraph
  • Update §6.4 with a precise statement of what variance evidence is and isn't reported
"""
import shutil, math, sys
from copy import deepcopy
from pathlib import Path

import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TS = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")
SRC = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v8_discussion_{TS}.docx"
DST = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v9_stats_{TS}.docx"
FIG_DIR = TRAIN_DIR / f"report_resources/figures/appendix_tb_plots/{TS}"


# ---------------------------------------------------------------
# Per-seed multi-seed measurements (verified earlier from BEST lines)
# ---------------------------------------------------------------
ADAPTED = {
    5: {"seed42": (0.006214, 0.8003),
        "seed1":  (0.006202, 0.8092),
        "seed2":  (0.006244, 0.8120)},
    6: {"seed42": (0.008876, 1.3294),
        "seed1":  (0.008927, 1.3316),
        "seed2":  (0.008891, 1.3277)},
    7: {"seed42": (0.009946, 1.7116),
        "seed1":  (0.013317, 2.5679),
        "seed2":  (0.013301, 2.5493)},
}
SINGLE_TASK = {5: (0.0093, 1.2039), 6: (0.0110, 1.7136), 7: (0.0101, 2.0853)}
SHARED      = {5: (0.0068, 0.9096), 6: (0.0092, 1.3689), 7: (0.0109, 1.9953)}

# t-critical for df = n-1 = 2 at 95% two-sided
T_975_DF2 = 4.302653
# t-critical for df = 2, one-sided p = 0.05
T_95_DF2_ONESIDED = 2.919986


def stats_for(seeds_dict):
    """Return (mean, std (ddof=1), 95% CI half-width via t-dist df=2)."""
    pos = np.array([v[0] for v in seeds_dict.values()])
    ori = np.array([v[1] for v in seeds_dict.values()])
    out = {}
    for name, arr in (("pos", pos), ("ori", ori)):
        m = float(arr.mean())
        s = float(arr.std(ddof=1))
        ci_half = T_975_DF2 * s / math.sqrt(len(arr))
        out[name] = {"mean": m, "std": s, "ci_low": m - ci_half, "ci_high": m + ci_half}
    return out


def one_sample_t_test_one_sided(sample, mu0, alternative="less"):
    """One-sample t-test. alternative='less' → H1: mean(sample) < mu0.

    Returns: t-statistic, df, p-value (one-sided), critical_t (one-sided 5%)."""
    arr = np.asarray(sample, dtype=float)
    n = len(arr); df = n - 1
    if df < 1:
        return None
    m = arr.mean(); s = arr.std(ddof=1)
    if s == 0:
        # Sample is degenerate; treat as deterministic
        if alternative == "less":
            return (-math.inf if m < mu0 else math.inf), df, 0.0 if m < mu0 else 1.0, T_95_DF2_ONESIDED
    se = s / math.sqrt(n)
    t = (m - mu0) / se
    # Approx one-sided p via t-distribution survival; we just report t and the critical
    return t, df, None, T_95_DF2_ONESIDED


def fmt_ci(stat):
    return f"{stat['mean']:.4f} ± {stat['std']:.4f} (95% CI: [{stat['ci_low']:.4f}, {stat['ci_high']:.4f}])"


def main():
    print("\n=== STATISTICAL ANALYSIS (n = 3 seeds, t-distribution df = 2) ===\n")
    rows = []   # for table-text
    test_summary = []

    for dof in (5, 6, 7):
        st = stats_for(ADAPTED[dof])
        rows.append((dof, st))
        print(f"--- DoF {dof} ---")
        print(f"  pos: {fmt_ci(st['pos'])} m")
        print(f"  ori: {fmt_ci(st['ori'])} °")

        # one-sample t-test vs SINGLE_TASK
        adp_pos = [v[0] for v in ADAPTED[dof].values()]
        adp_ori = [v[1] for v in ADAPTED[dof].values()]
        st_pos, st_ori = SINGLE_TASK[dof]
        sh_pos, sh_ori = SHARED[dof]

        t_pos_st, df_pos_st, _, t_crit = one_sample_t_test_one_sided(adp_pos, st_pos, "less")
        t_ori_st, df_ori_st, _, _ = one_sample_t_test_one_sided(adp_ori, st_ori, "less")
        t_pos_sh, _, _, _ = one_sample_t_test_one_sided(adp_pos, sh_pos, "less")
        t_ori_sh, _, _, _ = one_sample_t_test_one_sided(adp_ori, sh_ori, "less")

        def verdict(t):
            if t is None: return "n/a"
            if t == -math.inf: return "deterministically true"
            if t == math.inf: return "deterministically false"
            return f"t = {t:+.2f}, " + ("rejects H₀ (p < 0.05)" if t < -T_95_DF2_ONESIDED else "fails to reject H₀ (p ≥ 0.05)")

        print(f"  H₁: adapted < single-task")
        print(f"    pos: {verdict(t_pos_st)}")
        print(f"    ori: {verdict(t_ori_st)}")
        print(f"  H₁: adapted < shared")
        print(f"    pos: {verdict(t_pos_sh)}")
        print(f"    ori: {verdict(t_ori_sh)}")
        print()
        test_summary.append((dof, st, {
            "t_pos_st": t_pos_st, "t_ori_st": t_ori_st,
            "t_pos_sh": t_pos_sh, "t_ori_sh": t_ori_sh,
        }))

    # ---------------------------------------------------------------
    # Render updated Fig 5.7 with CI shading
    # ---------------------------------------------------------------
    fig, axes = plt.subplots(1, 2, figsize=(11.5, 4.6))
    ax_pos, ax_ori = axes
    dofs = [5, 6, 7]
    colours = {"seed42": "#1f77b4", "seed1": "#ff7f0e", "seed2": "#2ca02c"}
    markers = {"seed42": "o", "seed1": "s", "seed2": "^"}
    x_off = {"seed42": -0.10, "seed1": 0.0, "seed2": 0.10}

    # Per-seed scatter
    for seed in ("seed42", "seed1", "seed2"):
        for d in dofs:
            ax_pos.scatter(d + x_off[seed], ADAPTED[d][seed][0],
                           color=colours[seed], marker=markers[seed], s=70,
                           edgecolors="black", linewidths=0.6, zorder=4,
                           label=seed if d == 5 else None)
            ax_ori.scatter(d + x_off[seed], ADAPTED[d][seed][1],
                           color=colours[seed], marker=markers[seed], s=70,
                           edgecolors="black", linewidths=0.6, zorder=4,
                           label=seed if d == 5 else None)

    # Mean ± 95% CI (t-dist df=2)
    for ax, key in ((ax_pos, "pos"), (ax_ori, "ori")):
        means = []
        ci_low = []
        ci_high = []
        for d in dofs:
            st = stats_for(ADAPTED[d])
            means.append(st[key]["mean"])
            ci_low.append(st[key]["mean"] - st[key]["ci_low"])
            ci_high.append(st[key]["ci_high"] - st[key]["mean"])
        ax.errorbar(dofs, means, yerr=[ci_low, ci_high], fmt="D",
                    color="#9467bd", markersize=9, capsize=8,
                    elinewidth=1.6, label="mean (n=3) ± 95% CI", zorder=5)

    # Reference baselines
    for d in dofs:
        ax_pos.axhline(SINGLE_TASK[d][0], xmin=(d-4.5-0.6)/4.0, xmax=(d-4.5+0.6)/4.0,
                       color="#d62728", linestyle="--", linewidth=1.2, alpha=0.65)
        ax_ori.axhline(SINGLE_TASK[d][1], xmin=(d-4.5-0.6)/4.0, xmax=(d-4.5+0.6)/4.0,
                       color="#d62728", linestyle="--", linewidth=1.2, alpha=0.65)
        ax_pos.axhline(SHARED[d][0], xmin=(d-4.5-0.6)/4.0, xmax=(d-4.5+0.6)/4.0,
                       color="#7f7f7f", linestyle=":", linewidth=1.2, alpha=0.7)
        ax_ori.axhline(SHARED[d][1], xmin=(d-4.5-0.6)/4.0, xmax=(d-4.5+0.6)/4.0,
                       color="#7f7f7f", linestyle=":", linewidth=1.2, alpha=0.7)

    ax_pos.plot([], [], color="#d62728", linestyle="--", linewidth=1.2, label="single-task baseline")
    ax_pos.plot([], [], color="#7f7f7f", linestyle=":", linewidth=1.2, label="shared meta-kinematics")
    ax_ori.plot([], [], color="#d62728", linestyle="--", linewidth=1.2, label="single-task baseline")
    ax_ori.plot([], [], color="#7f7f7f", linestyle=":", linewidth=1.2, label="shared meta-kinematics")

    for ax in (ax_pos, ax_ori):
        ax.set_xticks(dofs); ax.set_xlabel("DoF configuration")
        ax.legend(loc="upper left", frameon=True, fontsize=9, ncol=1)
        ax.grid(True, linestyle="--", alpha=0.3)
    ax_pos.set_ylabel("Mean Euclidean position error (m)")
    ax_pos.set_title("(a) Adapted (best) position error across 3 seeds")
    ax_ori.set_ylabel("Mean orientation error (deg)")
    ax_ori.set_title("(b) Adapted (best) orientation error across 3 seeds")
    fig.suptitle("Multi-seed reproducibility with 95% CI (Student's t, df = 2)", fontsize=11)
    fig.tight_layout(rect=[0, 0, 1, 0.96])

    out_png = FIG_DIR / f"fig_multiseed_with_CI_{TS}.png"
    out_pdf = FIG_DIR / f"fig_multiseed_with_CI_{TS}.pdf"
    fig.savefig(out_png, dpi=200, bbox_inches="tight")
    fig.savefig(out_pdf, bbox_inches="tight")
    plt.close(fig)
    print(f"Rendered: {out_png}")

    # ---------------------------------------------------------------
    # Apply to docx
    # ---------------------------------------------------------------
    if not SRC.exists():
        sys.exit(f"Missing source: {SRC}")
    shutil.copy(str(SRC), str(DST))
    doc = docx.Document(str(DST))

    # ---------------------------------------------------------------
    # Update Table 5.1 cells: add 95% CI in addition to mean ± std.
    # We KEEP "mean ± std" form (don't overcrowd), but update caption.
    # ---------------------------------------------------------------
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Table 5.1") and "mean ± standard deviation" in p.text:
            new = (
                "Table 5.1  Comparison of forward-kinematics accuracy and training time across single-task, shared meta-kinematics "
                "and adapted meta-kinematics models on the KUKA iiwa 14. The Adapted (best) entries are reported as mean ± standard deviation "
                "over n = 3 random seeds (seeds 42, 1 and 2), with each seed varying parameter initialisation, support/query split and minibatch sampling "
                "while holding the Stage-2 shared checkpoint and hyperparameters fixed. Two-sided 95% confidence intervals "
                "(Student's t-distribution, df = 2) are reported in §5.4 below. Position errors are mean Euclidean errors in metres; "
                "orientation errors are mean geodesic angles in degrees."
            )
            runs = p.runs
            if runs:
                runs[0].text = new
                for r in runs[1:]:
                    r._element.getparent().remove(r._element)
            else:
                p.add_run(new)
            print("[OK] Updated Table 5.1 caption (mentions 95% CI)")
            break

    # ---------------------------------------------------------------
    # Replace the §5.4 fourth-observation paragraph with a more rigorous version
    # ---------------------------------------------------------------
    s5_pos = stats_for(ADAPTED[5])
    s6_pos = stats_for(ADAPTED[6])
    s7_pos = stats_for(ADAPTED[7])
    s5_ori = stats_for(ADAPTED[5])
    s6_ori = stats_for(ADAPTED[6])
    s7_ori = stats_for(ADAPTED[7])

    new_54 = (
        "A fourth observation comes from the multi-seed evaluation of the Adapted (best) row reported in Table 5.1. "
        f"For 5-DoF, the n = 3 sample mean position error is {s5_pos['pos']['mean']:.4f} m with sample standard deviation {s5_pos['pos']['std']:.5f} m and a Student's t-distribution 95% confidence interval of "
        f"[{s5_pos['pos']['ci_low']:.4f}, {s5_pos['pos']['ci_high']:.4f}] m (df = 2); orientation reaches {s5_ori['ori']['mean']:.3f}° ± {s5_ori['ori']['std']:.3f}°, 95% CI [{s5_ori['ori']['ci_low']:.3f}, {s5_ori['ori']['ci_high']:.3f}]°. "
        f"For 6-DoF the corresponding numbers are {s6_pos['pos']['mean']:.4f} ± {s6_pos['pos']['std']:.5f} m, 95% CI [{s6_pos['pos']['ci_low']:.4f}, {s6_pos['pos']['ci_high']:.4f}] m, and {s6_ori['ori']['mean']:.3f}° ± {s6_ori['ori']['std']:.3f}°, 95% CI [{s6_ori['ori']['ci_low']:.3f}, {s6_ori['ori']['ci_high']:.3f}]°. "
        "Across both these configurations the standard deviation is well below 1% of the mean, and the 95% CI does not overlap the corresponding single-task baseline; "
        "an unpaired one-sample Student's t-test (df = 2) confirms that the adapted model is statistically distinguishable from the single-task baseline at the p < 0.05 level on both metrics for the 5-DoF and 6-DoF configurations. "
        f"For 7-DoF the picture is materially different: the n = 3 mean position error is {s7_pos['pos']['mean']:.4f} ± {s7_pos['pos']['std']:.4f} m (95% CI [{s7_pos['pos']['ci_low']:.4f}, {s7_pos['pos']['ci_high']:.4f}] m), "
        f"and the orientation error is {s7_ori['ori']['mean']:.3f}° ± {s7_ori['ori']['std']:.3f}° (95% CI [{s7_ori['ori']['ci_low']:.3f}, {s7_ori['ori']['ci_high']:.3f}]°). "
        "These intervals encompass the single-task baseline (0.0101 m, 2.0853°) on the position metric and bracket it on the orientation metric, "
        "so under the present multi-seed protocol the 7-DoF adapted model cannot be claimed to outperform the single-task baseline at the 5% significance level — "
        "seed 42 attains a substantially better outcome than seeds 1 and 2, but with n = 3 the resulting CI is wide enough to include the baseline. "
        "This is the cleanest motivation for the multi-seed evaluation across all three stages flagged in Section 7.3 as the first item of the planned journal extension."
    )
    found = False
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("A fourth observation comes from the multi-seed evaluation"):
            runs = p.runs
            if runs:
                runs[0].text = new_54
                for r in runs[1:]:
                    r._element.getparent().remove(r._element)
            else:
                p.add_run(new_54)
            print("[OK] Replaced §5.4 fourth-observation paragraph with t-distribution CI version")
            found = True
            break
    if not found:
        print("[WARN] §5.4 fourth-observation paragraph not found")

    # ---------------------------------------------------------------
    # Replace the existing Fig 5.7 caption to mention CI / t-distribution
    # And replace the embedded picture with the new CI version
    # ---------------------------------------------------------------
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Figure 5.7") and "Multi-seed reproducibility" in p.text:
            new_cap = (
                "Figure 5.7  Multi-seed reproducibility of the per-DoF adaptation result. "
                "Three seeds (42, 1, 2) hold the Stage-2 shared meta-kinematics checkpoint and Stage-3 hyperparameters fixed and vary "
                "only parameter initialisation, support/query split and minibatch sampling. "
                "Filled markers show individual seed results; the diamond with whiskers is the n = 3 mean and the corresponding 95% confidence "
                "interval computed from a Student's t-distribution with df = 2. The dashed red line marks the single-task baseline and the "
                "dotted grey line the shared meta-kinematics value, both for the corresponding DoF. "
                "5-DoF and 6-DoF reproducibility is within the marker thickness; 7-DoF has notable seed sensitivity, with seed 42 attaining a "
                "lower error than seeds 1 and 2."
            )
            runs = p.runs
            if runs:
                runs[0].text = new_cap
                for r in runs[1:]:
                    r._element.getparent().remove(r._element)
            else:
                p.add_run(new_cap)
            print("[OK] Updated Figure 5.7 caption (mentions 95% CI)")
            break

    # Replace the actual Fig 5.7 image: there's no clean way to "replace" inline shapes
    # via python-docx, so we'll let the v9 docx keep the original Fig 5.7 image
    # and rely on the caption update. The new CI-shaded version is saved as
    # fig_multiseed_with_CI_{TS}.png in the figures dir for the user to drop in
    # via Word's Picture > Replace > Image option (10 seconds).
    print(f"[NOTE] New CI-shaded figure saved at: {FIG_DIR / f'fig_multiseed_with_CI_{TS}.png'}")
    print(f"       To replace the v9 Fig 5.7 image with this CI-shaded version,")
    print(f"       open the docx in Word, right-click the existing Fig 5.7 image,")
    print(f"       choose 'Change Picture' and select the CI-shaded PNG.")

    # ---------------------------------------------------------------
    # Replace §6.4 multi-seed paragraph with a more precise variance summary
    # ---------------------------------------------------------------
    for i, p in enumerate(doc.paragraphs):
        if "Adapted (best) row of Table 5.1 has been re-run for two further random seeds" in p.text:
            new_64 = (
                "Fourth, the statistical resolution of the comparisons in Chapter 5 is limited by the seed budget. "
                "The single-task and shared-meta-kinematics rows of Table 5.1 use a single fixed seed (42), which controls parameter initialisation, "
                "dataset shuffling and minibatch sampling, and is sufficient to characterise the qualitative pattern of results but does not provide a measure of seed-level variance for those two stages. "
                "The Adapted (best) row of Table 5.1 was re-run for two further random seeds (seeds 1 and 2) under the same Stage-2 checkpoint, "
                "hyperparameters and held-out split per configuration; the mean ± standard deviation reported there is therefore an n = 3 sample over the adaptation procedure, "
                "and the 95% confidence intervals computed in §5.4 from a Student's t-distribution with df = 2 reflect the genuine spread observed across the three seeds. "
                "These intervals capture variability in adaptation only, not variability that would arise from re-training Stages 1 and 2 with different seeds; the within-stage variance for those upstream stages is therefore unknown in this submission. "
                "The 7-DoF orientation gain reported in §5.3 of the original single-seed protocol is the most consequential improvement to test for full statistical robustness, and the n = 3 multi-seed evaluation has shown that this gain is partly driven by seed-42 specifically reaching a more favourable basin of attraction; "
                "a multi-seed evaluation of all three stages with at least n = 5 seeds, paired statistical tests across matched seeds, and held-out bootstrap 95% CIs is therefore the first item of the planned journal extension in Section 7.3, "
                "alongside the warm-start ablation (§5.6.1) and the head-to-head numerical comparison against SE(3)-aware architectures (§6.3) deferred to that journal version."
            )
            runs = p.runs
            if runs:
                runs[0].text = new_64
                for r in runs[1:]:
                    r._element.getparent().remove(r._element)
            else:
                p.add_run(new_64)
            print("[OK] Updated §6.4 multi-seed paragraph (precise variance summary)")
            break

    doc.save(str(DST))
    print(f"\n[OK] Saved: {DST}")
    print("\nSTATISTICAL TESTS SUMMARY:")
    print(f"  t-critical (df=2, one-sided p=0.05) = {T_95_DF2_ONESIDED:.3f}")
    print(f"  Reject H₀ if t < -{T_95_DF2_ONESIDED:.3f}")
    for dof, st, tests in test_summary:
        print(f"\n  DoF {dof}:")
        for k, v in tests.items():
            if v is None:
                continue
            if v == -math.inf:
                print(f"    {k}: t = -inf (deterministically passes)")
            elif v == math.inf:
                print(f"    {k}: t = +inf (deterministically fails)")
            else:
                ok = "✓ p<0.05" if v < -T_95_DF2_ONESIDED else "✗ p≥0.05"
                print(f"    {k}: t = {v:+.3f}  [{ok}]")


if __name__ == "__main__":
    main()
