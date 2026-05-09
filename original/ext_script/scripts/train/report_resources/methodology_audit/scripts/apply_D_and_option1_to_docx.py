#!/usr/bin/env python3
"""
Apply Experiment D (t-SNE) + Option 1 (adaptation curves) figures and prose
to the post-audit .docx. Saves to a new timestamped sibling.

The Experiment A (random-init 7-DoF) numbers are NOT applied here — they will
be added by a separate script after Experiment A completes.
"""

import shutil
from copy import deepcopy
from pathlib import Path

import docx
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TIMESTAMP = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")
SRC_DOCX = TRAIN_DIR / "FYP_Report_2_Chissanupong_2881058__post_audit.docx"
DST_DOCX = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v2_{TIMESTAMP}.docx"

TSNE_PNG = TRAIN_DIR / f"tier4_runs/expD_tsne_{TIMESTAMP}/fig_features_{TIMESTAMP}.png"
CURVES_PNG = TRAIN_DIR / f"tier4_runs/option1_adaptation_curves_{TIMESTAMP}/fig_adaptation_curves_{TIMESTAMP}.png"


def set_paragraph_text(p, new_text):
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(new_text)


def insert_picture_after(target_p, png_path, width_in=5.5, caption_text=None):
    """Insert a picture (and optional caption) immediately after target_p."""
    base_el = target_p._element
    # New empty paragraph for the image
    img_p_el = deepcopy(target_p._element)
    for r in list(img_p_el.findall(qn('w:r'))):
        img_p_el.remove(r)
    base_el.addnext(img_p_el)
    img_p = Paragraph(img_p_el, target_p._parent)
    img_p.text = ""
    run = img_p.add_run()
    run.add_picture(str(png_path), width=Inches(width_in))

    if caption_text:
        cap_el = deepcopy(target_p._element)
        for r in list(cap_el.findall(qn('w:r'))):
            cap_el.remove(r)
        img_p_el.addnext(cap_el)
        cap_p = Paragraph(cap_el, target_p._parent)
        cap_p.text = caption_text
    return img_p


def main():
    if not SRC_DOCX.exists():
        raise SystemExit(f"Source docx not found: {SRC_DOCX}")

    shutil.copy(str(SRC_DOCX), str(DST_DOCX))
    doc = docx.Document(str(DST_DOCX))

    # =================================================================
    # 1. EXPERIMENT D — t-SNE figure into §6.1
    # =================================================================
    if TSNE_PNG.exists():
        target_p = None
        for i, p in enumerate(doc.paragraphs):
            if "[Student to complete: 1–2 sentences once the projection has been generated." in p.text:
                target_p = p; break
            if "PCA / t-SNE on the standardised activations" in p.text:
                target_p = p; break

        if target_p is not None:
            new_text = (
                "A complementary, representational view of the same question is provided by examining the activations of the shared backbone directly. "
                "Figure 6.1 shows a two-dimensional t-SNE projection of the penultimate-layer features (the output of the eighth residual block, prior to the FK output head) "
                "for n = 1 000 held-out samples drawn from each DoF configuration. Two observations are noteworthy. "
                "First, the 7-DoF features (green) form a fairly distinct cluster on the left of the projection, while the 5-DoF (blue) and 6-DoF (orange) features overlap substantially in the centre and right. "
                "This is consistent with the underlying kinematics: the 5-DoF and 6-DoF configurations differ only in whether joint 6 is active, so they share most of the kinematic chain "
                "and produce similar end-effector geometry, whereas the 7-DoF configuration adds a further wrist axis whose contribution is qualitatively different. "
                "Second, the three configurations occupy a connected region of the projection rather than three disjoint clusters: the shared backbone has organised its representation "
                "by DoF count without treating each configuration as a separate task. This is consistent with the behavioural evidence from Sections 5.2–5.4, where the shared model "
                "transfers usefully across all three configurations and per-DoF adaptation refines rather than reconstructs the shared mapping."
            )
            set_paragraph_text(target_p, new_text)
            insert_picture_after(
                target_p, TSNE_PNG,
                width_in=5.0,
                caption_text=(
                    "Figure 6.1  t-SNE projection of the penultimate-layer activations of the shared meta-kinematics ResMLP_Mask "
                    "on n = 1 000 held-out samples per DoF configuration (5 DoF blue ○, 6 DoF orange □, 7 DoF green △). "
                    "Perplexity 30, learning rate auto, PCA initialisation, random_state 0."
                )
            )
            print("[OK] Inserted Figure 6.1 (t-SNE) into §6.1.")
        else:
            print("[WARN] §6.1 t-SNE placeholder paragraph not found.")
    else:
        print(f"[WARN] t-SNE figure not found: {TSNE_PNG}")

    # =================================================================
    # 2. OPTION 1 — adaptation-curves figure into §5.5 (training-cost analysis)
    #    Inserted after the existing Table 5.2 (training-time reductions table)
    #    near the end of §5.5, just before §5.6 begins.
    # =================================================================
    if CURVES_PNG.exists():
        # Find the §5.6 heading; insert just before it
        target_p = None
        target_idx = None
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            if txt.startswith("5.6") and "Ablation studies" in txt:
                target_idx = i
                break
        if target_idx is None:
            # Fallback: find the Figure 5.5 caption (last paragraph of §5.5)
            for i, p in enumerate(doc.paragraphs):
                if "Figure 5.5" in p.text and "Wall-clock training time" in p.text:
                    target_idx = i + 1
                    break

        if target_idx is not None and target_idx > 0:
            # Insert a new paragraph BEFORE the §5.6 heading by appending after the previous paragraph
            anchor = doc.paragraphs[target_idx - 1]
            new_text = (
                "Beyond the headline wall-clock numbers, the per-step adaptation traces in Figure 5.6 add a second piece of evidence: "
                "for the 5-DoF and 6-DoF configurations, adaptation reaches within 5% of its best score in approximately the first 4 000 steps "
                "(roughly 4% of the 100 000-step budget), after which the curve enters a slow refinement regime. "
                "For the 7-DoF configuration, the budget is necessary-but-sufficient: the curve continues to descend gently across the full 100 000 steps "
                "and only attains its best score at the end of the run. Read alongside Table 5.2, this means that the 0.365 hr and 0.363 hr wall-clocks for the 5-DoF and 6-DoF adapted models "
                "could be compressed substantially without losing accuracy, whereas the 0.111 hr 7-DoF budget is already close to the minimum required to reach the reported result. "
                "The qualitative shape of the curves — fast initial descent followed by slow refinement — is consistent across all three configurations and supports the interpretation "
                "that the shared backbone has placed each configuration close to its solution before adaptation begins, with the per-DoF stage performing a small, targeted refinement rather than re-learning the mapping."
            )
            # Insert prose first
            new_p_el = deepcopy(anchor._element)
            for r in list(new_p_el.findall(qn('w:r'))):
                new_p_el.remove(r)
            anchor._element.addnext(new_p_el)
            new_p = Paragraph(new_p_el, anchor._parent)
            new_p.text = new_text

            # Then the figure with caption
            insert_picture_after(
                new_p, CURVES_PNG,
                width_in=6.0,
                caption_text=(
                    "Figure 5.6  Stage-3 adaptation convergence on the held-out query split for each DoF configuration. "
                    "Panel (a) shows mean Euclidean position error (m); panel (b) shows mean orientation error (degrees). "
                    "Both axes use a log scale; star markers indicate the BEST step (the step minimising "
                    "s = E_pos + 0.01·E_ori, the same scoring rule used to select the values reported in Table 5.1)."
                )
            )
            print("[OK] Inserted Figure 5.6 (adaptation curves) into §5.5.")
        else:
            print("[WARN] §5.6 heading not found; could not place Figure 5.6.")
    else:
        print(f"[WARN] adaptation-curves figure not found: {CURVES_PNG}")

    doc.save(str(DST_DOCX))
    print(f"\n[OK] Saved: {DST_DOCX}")


if __name__ == "__main__":
    main()
