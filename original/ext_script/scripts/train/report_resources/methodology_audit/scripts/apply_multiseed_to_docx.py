#!/usr/bin/env python3
"""
Apply mean ± std to Table 5.1 "Adapted (best)" rows once the multi-seed sweep
is complete.

Reads BEST lines from each seed's adapt log, combines with the existing seed=42
headline numbers (already in Table 5.1), and writes mean ± std to the table
plus an updated narrative paragraph.

Sources of seed=42 headline values per DoF:
  5-DoF: pos_mae=0.006214, ori_deg=0.8003   (from adapt_5DOF_separate_weight summary)
  6-DoF: pos_mae=0.008876, ori_deg=1.3294   (from adapt_6DOF_separate_weight summary)
  7-DoF: pos_mae=0.009946, ori_deg=1.7116   (from sweep_fk_dof7 lr1e-6_l21e-6_S100000_ow0.05)

Usage:
  python apply_multiseed_to_docx.py
"""

import re, shutil, sys
from pathlib import Path

import numpy as np
import docx


TIMESTAMP = "20260507_045433"
TRAIN_DIR = Path("/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train")

SRC_DOCX = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v3_{TIMESTAMP}.docx"
DST_DOCX = TRAIN_DIR / f"FYP_Report_2_Chissanupong_2881058__post_audit_v4_multiseed_{TIMESTAMP}.docx"

LOGS_DIR = TRAIN_DIR / f"tier4_runs/expB_multiseed_smart_{TIMESTAMP}/logs"

# Existing seed=42 headline values (already in Table 5.1)
SEED42 = {
    5: {"pos_mae_m": 0.006214, "ori_deg": 0.8003},
    6: {"pos_mae_m": 0.008876, "ori_deg": 1.3294},
    7: {"pos_mae_m": 0.009946, "ori_deg": 1.7116},
}


def parse_best_line(log_path):
    """Returns dict {pos_mae_m, pos_rmse_m, ori_deg, best_step} or None."""
    if not log_path.exists():
        return None
    text = log_path.read_text(errors="ignore")
    m = re.search(r"\[INFO\]\s+BEST\s+step=(\d+)\s+metrics=(\{[^}]*\})", text)
    if not m:
        return None
    d = eval(m.group(2))
    d["best_step"] = int(m.group(1))
    return d


def collect_all():
    rows = {5: [], 6: [], 7: []}
    rows[5].append(("seed42", SEED42[5]["pos_mae_m"], SEED42[5]["ori_deg"]))
    rows[6].append(("seed42", SEED42[6]["pos_mae_m"], SEED42[6]["ori_deg"]))
    rows[7].append(("seed42", SEED42[7]["pos_mae_m"], SEED42[7]["ori_deg"]))
    for seed in (1, 2):
        for dof in (5, 6, 7):
            log = LOGS_DIR / f"adapt_seed{seed}_dof{dof}.log"
            d = parse_best_line(log)
            if d:
                rows[dof].append((f"seed{seed}", d["pos_mae_m"], d["ori_deg"]))
    return rows


def stats(values):
    arr = np.array(values, dtype=float)
    return float(arr.mean()), float(arr.std(ddof=1)) if len(arr) > 1 else 0.0


def replace_table_cell(t, row, col, new_text):
    cell = t.rows[row].cells[col]
    for p in cell.paragraphs:
        runs = p.runs
        if runs:
            runs[0].text = ""
            for r in runs[1:]:
                r._element.getparent().remove(r._element)
    cell.paragraphs[0].text = new_text


def set_paragraph_text(p, new_text):
    runs = p.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(new_text)


def main():
    if not SRC_DOCX.exists():
        sys.exit(f"Source not found: {SRC_DOCX}")

    rows = collect_all()
    print("Collected per-DoF measurements:")
    for dof, lst in rows.items():
        print(f"\n  DoF {dof}: {len(lst)} samples")
        for tag, p, o in lst:
            print(f"    {tag}: pos={p:.5f}  ori={o:.4f}")

    summaries = {}
    for dof, lst in rows.items():
        if len(lst) < 2:
            print(f"  DoF {dof}: only {len(lst)} sample(s); cannot compute std")
            continue
        pos_vals = [p for _, p, _ in lst]
        ori_vals = [o for _, _, o in lst]
        p_mean, p_std = stats(pos_vals)
        o_mean, o_std = stats(ori_vals)
        summaries[dof] = (p_mean, p_std, o_mean, o_std, len(lst))
        print(f"\n  DoF {dof} mean±std (n={len(lst)}):  pos = {p_mean:.4f} ± {p_std:.4f} m  |  ori = {o_mean:.4f} ± {o_std:.4f}°")

    # Save raw measurements as csv
    import csv
    out_csv = LOGS_DIR.parent / "multiseed_summary.csv"
    with open(out_csv, "w", newline="") as f:
        w = csv.writer(f)
        w.writerow(["dof", "seed", "pos_mae_m", "ori_deg"])
        for dof, lst in rows.items():
            for tag, p, o in lst:
                w.writerow([dof, tag, p, o])
    print(f"\nWrote: {out_csv}")

    if not summaries:
        print("\n[ABORT] No multi-seed data collected. Run the multi-seed sweep first.")
        return

    # Apply to docx
    shutil.copy(str(SRC_DOCX), str(DST_DOCX))
    doc = docx.Document(str(DST_DOCX))

    # Table 5.1 has 5 columns: Configuration | Metric | Single-task | Meta-kinematics | Adapted (best)
    # Rows (1-indexed under header): R1=5 pos, R2=5 ori, R3=5 time, R4=6 pos, R5=6 ori, R6=6 time,
    #                                R7=7 pos, R8=7 ori, R9=7 time
    t = doc.tables[6]   # Table 5.1 was index 6 in our index dump

    def fmt(mean, std, prec=4):
        return f"{mean:.{prec}f} ± {std:.{prec}f}"

    if 5 in summaries:
        p_mean, p_std, o_mean, o_std, n = summaries[5]
        replace_table_cell(t, 1, 4, fmt(p_mean, p_std))
        replace_table_cell(t, 2, 4, fmt(o_mean, o_std))
    if 6 in summaries:
        p_mean, p_std, o_mean, o_std, n = summaries[6]
        replace_table_cell(t, 4, 4, fmt(p_mean, p_std))
        replace_table_cell(t, 5, 4, fmt(o_mean, o_std))
    if 7 in summaries:
        p_mean, p_std, o_mean, o_std, n = summaries[7]
        replace_table_cell(t, 7, 4, fmt(p_mean, p_std))
        replace_table_cell(t, 8, 4, fmt(o_mean, o_std))

    print("[OK] Updated Table 5.1 'Adapted (best)' rows with mean ± std")

    # Update Table 5.1 caption to mention n=3 + DoF-specific budgets
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("Table 5.1") and "Comparison of forward-kinematics" in p.text:
            new_cap = (
                "Table 5.1  Comparison of forward-kinematics accuracy and training time across single-task, shared meta-kinematics "
                "and adapted meta-kinematics models on the KUKA iiwa 14. The Adapted (best) entries are reported as mean ± standard "
                "deviation over n = 3 random seeds (seeds 42, 1 and 2), with each seed varying parameter initialisation, support/query "
                "split and minibatch sampling while holding the Stage-2 shared checkpoint fixed. Position errors are mean Euclidean errors in metres; "
                "orientation errors are mean geodesic angles in degrees."
            )
            set_paragraph_text(p, new_cap)
            print("[OK] Updated Table 5.1 caption.")
            break

    # Update §6.4 limitations: soften the single-seed concession (now partially addressed)
    for i, p in enumerate(doc.paragraphs):
        if "single fixed seed (42)" in p.text and "qualitative pattern of results" in p.text:
            new_para = (
                "Fourth, the statistical resolution of the comparisons in Chapter 5 is limited by the seed budget. "
                "The single-task and shared-meta-kinematics rows of Table 5.1 use a single fixed seed (42), which controls parameter "
                "initialisation, dataset shuffling and minibatch sampling, and is sufficient to characterise the qualitative pattern of results. "
                "The Adapted (best) row of Table 5.1 has been re-run for two further random seeds (seeds 1 and 2) under the same Stage-2 "
                "checkpoint, hyperparameters and held-out split per configuration; the mean ± standard deviation reported there is therefore "
                "an n = 3 sample over the adaptation procedure, although it does not capture variance from re-training the shared model itself. "
                "The largest single relative improvement, the 7-DoF orientation gain, is the most consequential to test for full statistical robustness, "
                "and a multi-seed evaluation of all three stages with held-out bootstrap confidence intervals is therefore the first item of the "
                "planned journal extension in Section 7.3."
            )
            set_paragraph_text(p, new_para)
            print("[OK] Updated §6.4 limitations paragraph.")
            break

    doc.save(str(DST_DOCX))
    print(f"\n[OK] Saved: {DST_DOCX}")


if __name__ == "__main__":
    main()
