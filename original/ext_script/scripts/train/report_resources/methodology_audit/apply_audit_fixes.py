#!/usr/bin/env python3
"""
Apply Findings 1–16 from AUDIT_FINDINGS.md to the FYP .docx.

Strategy:
- Open the original .docx (read-only).
- Apply edits in-memory using python-docx + raw OOXML for inserts.
- Save to a separate file `__post_audit.docx`.
- The original file is NEVER modified by this script.
"""

from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = "/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train/FYP_Report_2_Chissanupong_2881058.docx"
DST = "/home/ubuntu/wish/kuka_kinematic_learner/original/ext_script/scripts/train/FYP_Report_2_Chissanupong_2881058__post_audit.docx"


def set_paragraph_text(p, new_text):
    """Replace all text in the paragraph with new_text. Keeps paragraph style.
    Drops formatting of individual runs (acceptable for prose)."""
    # Strip every run except the first; put new text in run 0
    runs = p.runs
    if runs:
        runs[0].text = new_text
        # Remove subsequent runs by deleting their XML elements
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        p.add_run(new_text)


def insert_paragraph_after(target_p, text, italic=False, copy_format_of=None):
    """Insert a new paragraph immediately after target_p. Returns the new Paragraph wrapper."""
    base = target_p._element
    # Build a new <w:p> by cloning the structure of a known paragraph (preserves pPr if we clone target)
    if copy_format_of is None:
        copy_format_of = target_p
    new_p_el = deepcopy(copy_format_of._element)
    # Strip all runs from the copy
    for r in list(new_p_el.findall(qn('w:r'))):
        new_p_el.remove(r)
    # Add a single new run
    r_el = OxmlElement('w:r')
    if italic:
        rPr = OxmlElement('w:rPr')
        i_el = OxmlElement('w:i')
        rPr.append(i_el)
        r_el.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = text
    t_el.set(qn('xml:space'), 'preserve')
    r_el.append(t_el)
    new_p_el.append(r_el)
    base.addnext(new_p_el)
    # Wrap in a Paragraph object
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p_el, target_p._parent)


def replace_table_cell(t, row, col, new_text):
    cell = t.rows[row].cells[col]
    # Clear all paragraphs in cell
    for p in cell.paragraphs:
        set_paragraph_text(p, "")
    # Put new text in first paragraph
    cell.paragraphs[0].text = new_text


def replace_listing_table(t, new_lines):
    """A code listing is a 1×1 table with a single multi-paragraph cell.
    Replace its content with new_lines, one paragraph per line."""
    cell = t.rows[0].cells[0]
    # Keep the first paragraph; clear it
    paragraphs = cell.paragraphs
    if not paragraphs:
        return
    first = paragraphs[0]
    set_paragraph_text(first, new_lines[0] if new_lines else "")
    # Remove subsequent paragraphs
    for p in list(paragraphs[1:]):
        p._element.getparent().remove(p._element)
    # Append the rest of the lines as new paragraphs in the cell
    for line in new_lines[1:]:
        new_p = cell.add_paragraph(line)


def main():
    doc = Document(SRC)

    # ============================================================
    # FINDING 4 + 5 + 7 + 8 — Paragraph 176 (§3.3 backbone description)
    # ============================================================
    new_para_176 = (
        "Each residual block applies a linear projection, a ReLU activation, dropout (p = 0.1) and a second linear projection, "
        "with the result added to the block's input through a skip connection followed by a final ReLU. "
        "The output head produces seven values corresponding to the predicted position t-hat ∈ R^3 and the predicted quaternion r-hat ∈ R^4; "
        "the quaternion is unit-normalised before any geodesic-distance computation in the loss and the evaluation metrics, "
        "but the model output itself is not normalised in-graph. "
        "In addition to the clamped joint vector, the model also receives the binary active-joint mask m_τ ∈ {0,1}^7 as an explicit auxiliary input. "
        "The mask is projected into the hidden space by a learned linear layer with no bias and added to the input projection of the joint vector, "
        "so that the model can adapt its internal representation to the active-DoF count rather than relying solely on the zero-clamped joint values. "
        "The mask projection is initialised to zero, which preserves backward compatibility with non-mask-aware single-task checkpoints "
        "when they are used to warm-start the shared meta-kinematics model in Stage 2. "
        "Inputs and targets are standardised per task before training, with each task's joint and pose statistics computed from its own training data "
        "and stored alongside the checkpoint. "
        "Stages 2 and 3 use one and the same residual MLP backbone, denoted ResMLP_Mask; Stage 3 fine-tunes the parameters that Stage 2 has produced. "
        "Stage 1's per-configuration backbone shares the same residual-block stack but does not include the mask projection, "
        "because each Stage-1 model only ever sees one DoF setting and does not need to be conditioned on it. "
        "Stage-1 checkpoints are loaded into the Stage-2 backbone with strict=False: the residual-block weights are inherited and the mask projection "
        "(which has no Stage-1 counterpart) is initialised to zero so that the Stage-2 model behaves identically to the warm-start Stage-1 model "
        "at the very first training step. This keeps the architecture simple while still supporting the meta-kinematics formulation of Section 3.1."
    )
    set_paragraph_text(doc.paragraphs[176], new_para_176)

    # ============================================================
    # FINDING 2 — Eq 3.3 split into 3.3a (Stages 1-2) and 3.3b (Stage 3)
    #   Original is paragraph 191 (the equation), 192 (the prose)
    # ============================================================
    set_paragraph_text(
        doc.paragraphs[191],
        "\tL^(1,2)(θ) = E_(q,p) [ ‖ f_θ(q) − p ‖² ],\t(3.3a)"
    )
    # Insert new equation paragraph for (3.3b) right after 191
    insert_paragraph_after(
        doc.paragraphs[191],
        "\tL^(3)(θ) = λ_p · E[ ‖ t-hat − t ‖² ] + λ_o · E[ θ_geo(r-hat, r)² ],\t(3.3b)"
    )

    # The prose paragraph following the equation now sits at 193 (it was 192 before our insert)
    new_para_after_eq = (
        "Equation (3.3a) describes the supervised objective used in Stages 1 and 2: a mean-squared error on the per-task standardised concatenated pose vector "
        "p = [t, r], minimised over the joint distribution of joint configurations and end-effector poses for the active task. "
        "Stage 1 uses a single task (K = 1); Stage 2 samples one of the K = 3 DoF configurations uniformly at each optimisation step. "
        "Equation (3.3b) describes the Stage-3 adaptation objective, which operates in raw physical units rather than standardised space and weights the "
        "position and orientation components explicitly. Here t and r are the un-standardised position and quaternion targets, t-hat and r-hat the corresponding "
        "model outputs, and θ_geo(r-hat, r) = 2 · arccos(|⟨r-hat, r⟩|) is the geodesic angle between the unit-normalised quaternions in radians. "
        "The split of the adaptation loss into raw-unit position MSE and squared geodesic angle is what allows the relative weight of position and orientation "
        "to be controlled in metres-and-degrees terms during the short fine-tuning stage."
    )
    # After our insert, original paragraph 192 has shifted to 193
    set_paragraph_text(doc.paragraphs[193], new_para_after_eq)

    # ============================================================
    # FINDING 9 + 8 — §4.4 software stack: rewrite held-out claim and add per-task standardisation
    # paragraph 206 was the original index; after our insert, it is now 207
    # ============================================================
    new_para_44 = (
        "All models are implemented in PyTorch [14] using the ResMLP backbone described in Chapter 3. "
        "Standardisation is performed per task: each DoF configuration's joint and pose statistics are computed from its own training data and stored "
        "alongside the checkpoint so that downstream evaluation and adaptation re-use the same normalisation. "
        "Inside Stage 2, training progress is monitored via a small quick-eval drawn from each task's data generator at every evaluation interval; "
        "the model with the lowest mean of the three quick-eval losses is preserved as the best checkpoint. "
        "Held-out test metrics reported in Chapter 5 are computed outside the training script, on a deterministic permutation subset of each DoF's dataset "
        "that is disjoint from the support set used for adaptation; the same permutation seed is used at evaluation time for all three model variants, "
        "so the held-out comparisons in Table 5.1 are over the same physical samples for the single-task, shared and adapted models in each configuration. "
        "Optimisation is performed with the Adam algorithm [16]. The shared meta-kinematics model is trained once across all three DoF settings, "
        "and per-DoF adaptation is then run separately for each target configuration and compared against both the shared model and the single-task baselines. "
        "Training is run on GPU within the Isaac Lab Python environment, so that data generation and model training share a single tensor-based pipeline. "
        "Specific hyperparameter values used for each stage are recorded in Appendix E and in the project repository; the architectural and protocol choices "
        "listed here are sufficient to reproduce the qualitative behaviour reported in Chapter 5."
    )
    set_paragraph_text(doc.paragraphs[207], new_para_44)

    # ============================================================
    # FINDING 1 — §4.5 / Eq 4.1 (RMSE -> mean Euclidean error)
    # Original 208/209/210 -> shifted to 209/210/211
    # ============================================================
    set_paragraph_text(
        doc.paragraphs[209],
        "Evaluation uses two pose-error metrics on the held-out split of each configuration's dataset. "
        "The position error is reported as the mean Euclidean error in metres between the predicted and ground-truth end-effector positions, "
        "which has a direct physical interpretation as the average length of the position-error vector across the held-out split:"
    )
    set_paragraph_text(
        doc.paragraphs[210],
        "\tE_pos = (1/N) Σᵢ ‖ t-hatᵢ − tᵢ ‖₂,\t(4.1)"
    )
    # The next paragraph (originally 210) introduces orientation error; extend with the explicit formula.
    new_para_ori = (
        "The orientation error is the mean geodesic angle between the predicted and ground-truth quaternions, reported in degrees so that errors are "
        "interpretable on the same physical scale used in industrial robot specifications: "
        "E_ori = (1/N) Σᵢ 2 · arccos(|⟨ r-hatᵢ, rᵢ ⟩|) · (180/π). "
        "In addition to the two error metrics, wall-clock training time is recorded for every model and every stage, in order to compare the cost of "
        "training a per-configuration model from scratch with the much shorter cost of adapting from the shared meta-kinematics initialisation. "
        "All three metrics — mean position error in metres, mean orientation error in degrees and training time in hours — are reported jointly "
        "for the single-task, shared meta-kinematics and adapted meta-kinematics models in Chapter 5."
    )
    set_paragraph_text(doc.paragraphs[211], new_para_ori)

    # ============================================================
    # FINDING 5 — §4.6 backbone-sharing wording (original para 213 -> 214 after insert)
    # ============================================================
    # Paragraph at index 214 should be the §4.6 second paragraph.
    # Let's grep for the marker to be safe — but we know our offset is +1 after the Eq-3.3b insert.
    target_idx = 214  # original 213 + 1
    p = doc.paragraphs[target_idx]
    if "share the same ResMLP backbone" not in p.text:
        # Try +/- 1 to find the right paragraph
        for cand in (213, 214, 215, 212, 216):
            if "share the same ResMLP backbone" in doc.paragraphs[cand].text:
                target_idx = cand
                break
    if "share the same ResMLP backbone" in doc.paragraphs[target_idx].text:
        new_46 = (
            "To make the comparison across the three model variants fair, a number of confounders are controlled. "
            "The three variants — single-task, shared meta-kinematics and per-DoF adapted — share the same residual-block stack and the same FK output head; "
            "Stages 2 and 3 additionally include the mask-conditioning projection introduced for the multi-DoF setting (Section 3.3). "
            "The same dataset and held-out split per configuration, the same optimiser family, and the same input-and-target standardisation are used across all three variants. "
            "The differences are limited to the training schedule, the loss form (a unified MSE on the standardised pose vector in Stages 1–2 versus a weighted raw-unit loss "
            "in Stage 3) and the initialisation of parameters. Where Stage 3 (adaptation) is contrasted with Stage 1 (single-task) on equal time budgets, "
            "the time budget is set to the wall-clock time of the longer Stage 1 run, so that the Stage 3 model cannot trivially benefit from running for less time than the comparison."
        )
        set_paragraph_text(doc.paragraphs[target_idx], new_46)

    # ============================================================
    # FINDING 10 — Disclose Stage-3 best-step scoring
    # Add a sentence to §5.3 (Per-DoF adaptation) — paragraph index after the §5.3 heading
    # ============================================================
    # Find the §5.3 paragraph that contains "Per-DoF adaptation consistently improves over both"
    for i, pp in enumerate(doc.paragraphs):
        if "Per-DoF adaptation consistently improves over both the shared" in pp.text:
            old = pp.text
            new = old + (
                " For each adapted run, the reported result is the evaluation step that minimises the scalarised score "
                "s = w_p · E_pos + w_o · E_ori with w_p = 1.0 and w_o = 0.01, which keeps position the dominant criterion "
                "while breaking ties using orientation; the same weights are used across the three DoF configurations to avoid cherry-picking a per-metric best step."
            )
            set_paragraph_text(pp, new)
            break

    # ============================================================
    # FINDING 4 (Listing E.2) — replace LayerNorm version with Dropout version
    # Listings are tables 17, 18, 19
    # ============================================================
    # Listing E.2: ResBlock
    listing_e2 = [
        "class ResBlock(nn.Module):",
        "    def __init__(self, dim: int = 1024, p_drop: float = 0.1):",
        "        super().__init__()",
        "        self.fc1  = nn.Linear(dim, dim)",
        "        self.fc2  = nn.Linear(dim, dim)",
        "        self.act  = nn.ReLU()",
        "        self.drop = nn.Dropout(p_drop)",
        "",
        "    def forward(self, x):",
        "        h = self.act(self.fc1(x))",
        "        h = self.drop(h)",
        "        h = self.fc2(h)",
        "        return self.act(h + x)        # ReLU after additive skip",
    ]
    replace_listing_table(doc.tables[18], listing_e2)

    # ============================================================
    # FINDING 2 (Listing E.1) — relabel as Stage-3 quaternion geodesic loss
    # ============================================================
    listing_e1 = [
        "def quat_geodesic_loss_rad2(r_hat, r, eps=1e-7):",
        "    # r_hat, r: (B, 4) tensors in (qw, qx, qy, qz) order.",
        "    # Sign-invariant geodesic angle (radians), squared and averaged.",
        "    r_hat = r_hat / r_hat.norm(dim=1, keepdim=True).clamp_min(eps)",
        "    r     = r     / r.norm(dim=1, keepdim=True).clamp_min(eps)",
        "    dot = (r_hat * r).sum(dim=1).abs().clamp(eps, 1.0 - eps)",
        "    ang = 2.0 * torch.acos(dot)        # in [0, pi]",
        "    return (ang ** 2).mean()           # used inside the Stage-3 loss (3.3b)",
    ]
    replace_listing_table(doc.tables[17], listing_e1)

    # ============================================================
    # FINDING 6 (Listing E.3) — uniform-task sampling, not K-task average
    # ============================================================
    listing_e3 = [
        "def multitask_step(model, task_loaders, optimizer, mse):",
        "    # One Stage-2 optimisation step on a uniformly-sampled DoF configuration.",
        "    tname = random.choice(list(task_loaders.keys()))   # 5dof / 6dof / 7dof",
        "    x, y, mask = next(task_loaders[tname])             # already per-task standardised",
        "    yhat = model(x, mask)",
        "    loss = mse(yhat, y)                                 # MSE on standardised pose",
        "    optimizer.zero_grad(); loss.backward(); optimizer.step()",
        "    return tname, loss.item()",
    ]
    replace_listing_table(doc.tables[19], listing_e3)

    # Update Listing E.1 caption ("Quaternion-distance term used inside the multi-task loss (3.3).")
    for i, pp in enumerate(doc.paragraphs):
        if pp.text.startswith("Listing E.1") and "Quaternion-distance" in pp.text:
            set_paragraph_text(pp, "Listing E.1  Squared geodesic-angle term used inside the Stage-3 adaptation loss (3.3b).")
            break

    # ============================================================
    # FINDING 3, 4, 13, 12, 14, 15 — Table E.1 (table index 15)
    # Rows of E.1 (18 rows × 2 cols), R0 = header
    # R6 was "LayerNorm (within each residual block)"
    # ============================================================
    e1 = doc.tables[15]
    # Print current rows for safety
    # Replace R6 (LayerNorm) with Dropout/regularisation
    replace_table_cell(e1, 5, 0, "Activation")  # R5 was Activation: ReLU; keep ReLU
    replace_table_cell(e1, 5, 1, "ReLU")
    replace_table_cell(e1, 6, 0, "Regularisation (each residual block)")
    replace_table_cell(e1, 6, 1, "Dropout (p = 0.1) between the two linear layers; no LayerNorm or BatchNorm")
    # R11 was Weight decay
    replace_table_cell(e1, 11, 0, "Weight decay")
    replace_table_cell(e1, 11, 1, "1×10⁻⁵ (Adam L2, Stage 1); 0 (Stage 2); 1×10⁻⁶ L2-to-init (Stage 3 adaptation)")
    # R12 = Inactive-joint encoding -> change to mask-conditioned + zero-clamp
    replace_table_cell(e1, 12, 0, "Mask conditioning")
    replace_table_cell(e1, 12, 1, "Inactive joints zero-clamped at the input; binary active-joint mask additionally projected into the hidden space by a learned, zero-initialised linear layer (Stages 2–3)")
    # R14 = std_floor_q — qualify it as IK-only
    replace_table_cell(e1, 14, 0, "Joint-noise floor (IK target standardisation)")
    replace_table_cell(e1, 14, 1, "std_floor_q = 1.0° equivalent. Not used in FK — the work reported here is FK only")

    # ============================================================
    # FINDING 11, 13, 16 — Table E.2 (table index 16)
    # ============================================================
    e2 = doc.tables[16]
    # Headers are R0
    # R1: Initial learning rate -> 5e-4 / 3e-4 / 1e-5
    replace_table_cell(e2, 1, 0, "Initial learning rate")
    replace_table_cell(e2, 1, 1, "5 × 10⁻⁴")
    replace_table_cell(e2, 1, 2, "3 × 10⁻⁴")
    replace_table_cell(e2, 1, 3, "1 × 10⁻⁵")
    # R2: LR schedule
    replace_table_cell(e2, 2, 0, "LR schedule")
    replace_table_cell(e2, 2, 1, "ReduceLROnPlateau (factor 0.5, patience 10)")
    replace_table_cell(e2, 2, 2, "Cosine, linear warmup of 2 000 steps")
    replace_table_cell(e2, 2, 3, "Constant")
    # R3: LR minimum
    replace_table_cell(e2, 3, 0, "LR minimum")
    replace_table_cell(e2, 3, 1, "— (set by plateau decay)")
    replace_table_cell(e2, 3, 2, "1 × 10⁻⁵")
    replace_table_cell(e2, 3, 3, "—")
    # R4: Batch size
    replace_table_cell(e2, 4, 0, "Batch size")
    replace_table_cell(e2, 4, 1, "8 192")
    replace_table_cell(e2, 4, 2, "4 096")
    replace_table_cell(e2, 4, 3, "8 192")
    # R5: Total steps
    replace_table_cell(e2, 5, 0, "Total steps")
    replace_table_cell(e2, 5, 1, "Up to 200 epochs (early-stopped on val loss plateau)")
    replace_table_cell(e2, 5, 2, "300 000 (best at step 289 500)")
    replace_table_cell(e2, 5, 3, "100 000")
    # R6: Support / query size — Stage 1 doesn't have it; Stage 3 yes
    replace_table_cell(e2, 6, 0, "Support / query size")
    replace_table_cell(e2, 6, 1, "n/a (50/20/30 train/val/test split)")
    replace_table_cell(e2, 6, 2, "n/a")
    replace_table_cell(e2, 6, 3, "50 k support / 2 M query")
    # R7: Gradient clipping
    replace_table_cell(e2, 7, 0, "Gradient clipping (max norm)")
    replace_table_cell(e2, 7, 1, "1.0")
    replace_table_cell(e2, 7, 2, "1.0")
    replace_table_cell(e2, 7, 3, "1.0")
    # R8: L2 / weight decay — split semantics across stages
    replace_table_cell(e2, 8, 0, "Weight decay / L2")
    replace_table_cell(e2, 8, 1, "1 × 10⁻⁵ (Adam L2)")
    replace_table_cell(e2, 8, 2, "0")
    replace_table_cell(e2, 8, 3, "1 × 10⁻⁶ (L2-to-init regulariser)")
    # R9: Wall-clock 5 DoF — keep numbers
    # R10: Wall-clock 6 DoF
    # R11: Wall-clock 7 DoF
    # R12: Random seed — already correct ("42") — leave alone

    # Save
    doc.save(DST)
    print(f"Saved post-audit report to: {DST}")


if __name__ == "__main__":
    main()
